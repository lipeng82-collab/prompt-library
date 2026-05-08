# -*- coding: utf-8 -*-
"""
湘小印·文档出品 - 版式转换引擎
将Markdown文件按照指定版式模板转换为Word文档（.docx）

支持的版式模板：
  - guowei    : 省国资委公文（默认）
  - xingxiang : 兴湘集团文件
  - neican    : 内参简报

用法：
  py -3 md_to_docx.py input.md output.docx
  py -3 md_to_docx.py input.md output.docx --template xingxiang
  py -3 md_to_docx.py input.md output.docx --template neican
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import argparse
import sys


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  版式模板定义
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

TEMPLATES = {
    "guowei": {
        "name": "省国资委公文",
        # 页面设置
        "page": {
            "width_cm": 21.0, "height_cm": 29.7,
            "top_cm": 2.25, "bottom_cm": 2.25,
            "left_cm": 3.25, "right_cm": 3.17,
        },
        # 字体设置（每种元素 = dict）
        "styles": {
            "h1": {"font_cn": "方正小标宋简体", "font_en": "Times New Roman", "size_pt": 22,
                   "bold": False, "align": "center", "first_line_cm": None,
                   "space_before_pt": 6, "space_after_pt": 12, "line_spacing_pt": None,
                   "line_spacing_rule": "single"},
            "h2": {"font_cn": "黑体", "font_en": "Times New Roman", "size_pt": 16,
                   "bold": False, "align": None, "first_line_cm": 1.13,
                   "space_before_pt": 6, "space_after_pt": 0, "line_spacing_pt": 28,
                   "line_spacing_rule": "exactly"},
            "h3": {"font_cn": "楷体_GB2312", "font_en": "Times New Roman", "size_pt": 16,
                   "bold": True, "align": None, "first_line_cm": 1.13,
                   "space_before_pt": 4, "space_after_pt": 0, "line_spacing_pt": 28,
                   "line_spacing_rule": "exactly"},
            "h4": {"font_cn": "楷体_GB2312", "font_en": "Times New Roman", "size_pt": 16,
                   "bold": True, "align": None, "first_line_cm": 1.13,
                   "space_before_pt": 2, "space_after_pt": 0, "line_spacing_pt": 28,
                   "line_spacing_rule": "exactly"},
            "body": {"font_cn": "仿宋_GB2312", "font_en": "Times New Roman", "size_pt": 16,
                     "bold": False, "align": None, "first_line_cm": 1.13,
                     "space_before_pt": 0, "space_after_pt": 0, "line_spacing_pt": 28,
                     "line_spacing_rule": "exactly"},
            "table": {"font_cn": "仿宋_GB2312", "font_en": "Times New Roman", "size_pt": 14,
                      "bold": False, "line_spacing_pt": 22},
            "footer": {"font_cn": "仿宋_GB2312", "font_en": "Times New Roman", "size_pt": 16,
                       "bold": False, "align": "right",
                       "space_before_pt": 0, "space_after_pt": 0, "line_spacing_pt": 28,
                       "line_spacing_rule": "exactly"},
        },
    },
    "xingxiang": {
        "name": "兴湘集团文件",
        # 页面设置（与省国资委公文一致）
        "page": {
            "width_cm": 21.0, "height_cm": 29.7,
            "top_cm": 2.25, "bottom_cm": 2.25,
            "left_cm": 3.25, "right_cm": 3.17,
        },
        "styles": {
            "h1": {"font_cn": "方正小标宋简体", "font_en": "Times New Roman", "size_pt": 22,
                   "bold": False, "align": "center", "first_line_cm": None,
                   "space_before_pt": 6, "space_after_pt": 12, "line_spacing_pt": None,
                   "line_spacing_rule": "single"},
            "h2": {"font_cn": "黑体", "font_en": "Times New Roman", "size_pt": 16,
                   "bold": False, "align": None, "first_line_cm": 1.13,
                   "space_before_pt": 6, "space_after_pt": 0, "line_spacing_pt": 28,
                   "line_spacing_rule": "exactly"},
            "h3": {"font_cn": "楷体_GB2312", "font_en": "Times New Roman", "size_pt": 16,
                   "bold": True, "align": None, "first_line_cm": 1.13,
                   "space_before_pt": 4, "space_after_pt": 0, "line_spacing_pt": 28,
                   "line_spacing_rule": "exactly"},
            "h4": {"font_cn": "楷体_GB2312", "font_en": "Times New Roman", "size_pt": 16,
                   "bold": True, "align": None, "first_line_cm": 1.13,
                   "space_before_pt": 2, "space_after_pt": 0, "line_spacing_pt": 28,
                   "line_spacing_rule": "exactly"},
            "body": {"font_cn": "仿宋_GB2312", "font_en": "Times New Roman", "size_pt": 16,
                     "bold": False, "align": None, "first_line_cm": 1.13,
                     "space_before_pt": 0, "space_after_pt": 0, "line_spacing_pt": 28,
                     "line_spacing_rule": "exactly"},
            "table": {"font_cn": "仿宋_GB2312", "font_en": "Times New Roman", "size_pt": 14,
                      "bold": False, "line_spacing_pt": 22},
            "footer": {"font_cn": "仿宋_GB2312", "font_en": "Times New Roman", "size_pt": 16,
                       "bold": False, "align": "right",
                       "space_before_pt": 0, "space_after_pt": 0, "line_spacing_pt": 28,
                       "line_spacing_rule": "exactly"},
        },
    },
    "neican": {
        "name": "内参简报",
        "page": {
            "width_cm": 21.0, "height_cm": 29.7,
            "top_cm": 2.54, "bottom_cm": 2.54,
            "left_cm": 3.17, "right_cm": 3.17,
        },
        "styles": {
            "h1": {"font_cn": "方正小标宋简体", "font_en": "Times New Roman", "size_pt": 20,
                   "bold": False, "align": "center", "first_line_cm": None,
                   "space_before_pt": 6, "space_after_pt": 10, "line_spacing_pt": None,
                   "line_spacing_rule": "single"},
            "h2": {"font_cn": "黑体", "font_en": "Times New Roman", "size_pt": 15,
                   "bold": False, "align": None, "first_line_cm": 0.74,
                   "space_before_pt": 6, "space_after_pt": 0, "line_spacing_pt": 28,
                   "line_spacing_rule": "exactly"},
            "h3": {"font_cn": "楷体_GB2312", "font_en": "Times New Roman", "size_pt": 15,
                   "bold": True, "align": None, "first_line_cm": 0.74,
                   "space_before_pt": 4, "space_after_pt": 0, "line_spacing_pt": 28,
                   "line_spacing_rule": "exactly"},
            "h4": {"font_cn": "楷体_GB2312", "font_en": "Times New Roman", "size_pt": 15,
                   "bold": True, "align": None, "first_line_cm": 0.74,
                   "space_before_pt": 2, "space_after_pt": 0, "line_spacing_pt": 28,
                   "line_spacing_rule": "exactly"},
            "body": {"font_cn": "仿宋_GB2312", "font_en": "Times New Roman", "size_pt": 15,
                     "bold": False, "align": None, "first_line_cm": 0.74,
                     "space_before_pt": 0, "space_after_pt": 0, "line_spacing_pt": 28,
                     "line_spacing_rule": "exactly"},
            "table": {"font_cn": "仿宋_GB2312", "font_en": "Times New Roman", "size_pt": 13,
                      "bold": False, "line_spacing_pt": 22},
            "footer": {"font_cn": "仿宋_GB2312", "font_en": "Times New Roman", "size_pt": 15,
                       "bold": False, "align": "right",
                       "space_before_pt": 0, "space_after_pt": 0, "line_spacing_pt": 28,
                       "line_spacing_rule": "exactly"},
        },
    },
}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  辅助函数
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def set_page_layout(doc, tmpl):
    """设置页面版式"""
    page = tmpl["page"]
    section = doc.sections[0]
    section.page_width  = Cm(page["width_cm"])
    section.page_height = Cm(page["height_cm"])
    section.top_margin    = Cm(page["top_cm"])
    section.bottom_margin = Cm(page["bottom_cm"])
    section.left_margin   = Cm(page["left_cm"])
    section.right_margin  = Cm(page["right_cm"])


def set_para_format(para, first_line_cm=None, left_indent_cm=None,
                    space_before_pt=0, space_after_pt=0,
                    line_spacing_pt=28, line_spacing_rule="exactly",
                    align=None):
    """设置段落格式"""
    pf = para.paragraph_format
    if first_line_cm is not None:
        pf.first_line_indent = Cm(first_line_cm)
    if left_indent_cm is not None:
        pf.left_indent = Cm(left_indent_cm)
    pf.space_before = Pt(space_before_pt)
    pf.space_after  = Pt(space_after_pt)
    if line_spacing_rule == "single":
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing_pt)
    if align is not None:
        para.alignment = align


def add_run(para, text, font_cn=None, font_en=None,
            size_pt=16, bold=False, italic=False, color=None):
    """添加run并设置字体"""
    run = para.add_run(text)
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_cn or font_en:
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        if font_cn:
            rFonts.set(qn('w:eastAsia'), font_cn)
        if font_en:
            rFonts.set(qn('w:ascii'),    font_en)
            rFonts.set(qn('w:hAnsi'),    font_en)
        if font_cn and not font_en:
            rFonts.set(qn('w:ascii'),    font_cn)
            rFonts.set(qn('w:hAnsi'),    font_cn)
    return run


def resolve_align(align_str):
    """将字符串对齐方式转换为docx枚举"""
    if align_str == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    elif align_str == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    elif align_str == "left":
        return WD_ALIGN_PARAGRAPH.LEFT
    return None


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  行类型识别
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# 编制信息识别关键词
_FOOTER_KEYWORDS = [
    "报告编制单位", "报告日期", "编制单位", "编制日期",
    "发布单位", "发布日期", "方案编制单位", "方案发布日期",
]


def classify_line(line):
    """返回 (type, content)"""
    stripped = line.strip()
    if not stripped or stripped in ('---', '---\n'):
        return ('empty', '')
    if stripped.startswith('# ') and not stripped.startswith('## '):
        return ('h1', stripped[2:].strip())
    if stripped.startswith('## ') and not stripped.startswith('### '):
        return ('h2', stripped[3:].strip())
    if stripped.startswith('### ') and not stripped.startswith('#### '):
        return ('h3', stripped[4:].strip())
    if stripped.startswith('#### '):
        return ('h4', stripped[5:].strip())
    if stripped.startswith('|'):
        return ('table_row', stripped)
    for kw in _FOOTER_KEYWORDS:
        if stripped.startswith('**' + kw):
            return ('footer_info', stripped)
    return ('body', stripped)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  富文本渲染
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def render_inline(para, text, font_cn, font_en, size_pt, base_bold=False):
    """解析 markdown **加粗** 语法，分段渲染到段落"""
    pattern = re.compile(r'\*\*(.+?)\*\*')
    pos = 0
    for m in pattern.finditer(text):
        before = text[pos:m.start()]
        if before:
            add_run(para, before, font_cn, font_en, size_pt, bold=base_bold)
        inner = m.group(1)
        add_run(para, inner, font_cn, font_en, size_pt, bold=True)
        pos = m.end()
    tail = text[pos:]
    if tail:
        add_run(para, tail, font_cn, font_en, size_pt, bold=base_bold)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  表格渲染
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def render_table(doc, table_lines, tmpl):
    """将 markdown 表格渲染为 word 表格"""
    rows_data = []
    for line in table_lines:
        if re.match(r'^\|[-| :]+\|$', line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows_data.append(cells)
    if not rows_data:
        return

    ts = tmpl["styles"]["table"]
    col_count = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=col_count)
    tbl.style = 'Table Grid'

    for ri, row_data in enumerate(rows_data):
        for ci, cell_text in enumerate(row_data):
            if ci >= col_count:
                break
            cell = tbl.cell(ri, ci)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after  = Pt(2)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(ts["line_spacing_pt"])
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            is_header = (ri == 0)
            add_run(p, clean,
                    font_cn=ts["font_cn"], font_en=ts["font_en"],
                    size_pt=ts["size_pt"], bold=is_header)
    doc.add_paragraph()


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  段落判断辅助
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def is_intro_para(text, prev_type):
    """判断是否为章节导语段（跟在 ## 或 ### 后面的首段文字，无特殊标记）"""
    if prev_type in ('h2', 'h3'):
        if not text.startswith('**') and not text.startswith('一是') \
           and not text.startswith('二是') and not text.startswith('三是') \
           and not text.startswith('四是') and not text.startswith('五是'):
            return True
    return False


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  核心转换函数
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def md_to_docx(md_path, out_path, template_name="guowei"):
    """
    将 Markdown 文件转换为 Word 文档

    参数:
        md_path        : 输入Markdown文件路径
        out_path       : 输出Word文件路径
        template_name  : 版式模板名称（guowei/xingxiang/neican）
    """
    if template_name not in TEMPLATES:
        print(f"错误：未知版式模板 '{template_name}'")
        print(f"可用模板：{', '.join(TEMPLATES.keys())}")
        sys.exit(1)

    tmpl = TEMPLATES[template_name]
    print(f"版式模板：{tmpl['name']}")

    doc = Document()
    set_page_layout(doc, tmpl)

    # 删除默认空段落
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    # 预处理：合并表格行
    processed = []
    i = 0
    while i < len(lines):
        t, c = classify_line(lines[i])
        if t == 'table_row':
            table_block = []
            while i < len(lines) and classify_line(lines[i])[0] == 'table_row':
                table_block.append(classify_line(lines[i])[1])
                i += 1
            processed.append(('table_block', table_block))
        else:
            processed.append((t, c))
            i += 1

    ss = tmpl["styles"]
    prev_type = None

    for (t, c) in processed:
        if t in ('empty',):
            prev_type = t
            continue

        if t == 'h1':
            s = ss["h1"]
            p = doc.add_paragraph()
            align = resolve_align(s.get("align"))
            set_para_format(p,
                first_line_cm=s.get("first_line_cm"),
                space_before_pt=s.get("space_before_pt", 0),
                space_after_pt=s.get("space_after_pt", 0),
                line_spacing_pt=s.get("line_spacing_pt", 28),
                line_spacing_rule=s.get("line_spacing_rule", "exactly"),
                align=align)
            add_run(p, c,
                    font_cn=s["font_cn"], font_en=s["font_en"],
                    size_pt=s["size_pt"], bold=s["bold"])

        elif t == 'h2':
            s = ss["h2"]
            p = doc.add_paragraph()
            set_para_format(p,
                first_line_cm=s.get("first_line_cm"),
                space_before_pt=s.get("space_before_pt", 0),
                space_after_pt=s.get("space_after_pt", 0),
                line_spacing_pt=s.get("line_spacing_pt", 28),
                line_spacing_rule=s.get("line_spacing_rule", "exactly"))
            add_run(p, c,
                    font_cn=s["font_cn"], font_en=s["font_en"],
                    size_pt=s["size_pt"], bold=s["bold"])

        elif t == 'h3':
            s = ss["h3"]
            p = doc.add_paragraph()
            set_para_format(p,
                first_line_cm=s.get("first_line_cm"),
                space_before_pt=s.get("space_before_pt", 0),
                space_after_pt=s.get("space_after_pt", 0),
                line_spacing_pt=s.get("line_spacing_pt", 28),
                line_spacing_rule=s.get("line_spacing_rule", "exactly"))
            add_run(p, c,
                    font_cn=s["font_cn"], font_en=s["font_en"],
                    size_pt=s["size_pt"], bold=s["bold"])

        elif t == 'h4':
            s = ss["h4"]
            p = doc.add_paragraph()
            set_para_format(p,
                first_line_cm=s.get("first_line_cm"),
                space_before_pt=s.get("space_before_pt", 0),
                space_after_pt=s.get("space_after_pt", 0),
                line_spacing_pt=s.get("line_spacing_pt", 28),
                line_spacing_rule=s.get("line_spacing_rule", "exactly"))
            add_run(p, c,
                    font_cn=s["font_cn"], font_en=s["font_en"],
                    size_pt=s["size_pt"], bold=s["bold"])

        elif t == 'table_block':
            render_table(doc, c, tmpl)

        elif t == 'footer_info':
            s = ss["footer"]
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', c)
            p = doc.add_paragraph()
            align = resolve_align(s.get("align"))
            set_para_format(p,
                first_line_cm=s.get("first_line_cm"),
                space_before_pt=s.get("space_before_pt", 0),
                space_after_pt=s.get("space_after_pt", 0),
                line_spacing_pt=s.get("line_spacing_pt", 28),
                line_spacing_rule=s.get("line_spacing_rule", "exactly"),
                align=align)
            add_run(p, clean,
                    font_cn=s["font_cn"], font_en=s["font_en"],
                    size_pt=s["size_pt"], bold=s["bold"])

        elif t == 'body':
            s = ss["body"]
            p = doc.add_paragraph()
            set_para_format(p,
                first_line_cm=s.get("first_line_cm"),
                space_before_pt=s.get("space_before_pt", 0),
                space_after_pt=s.get("space_after_pt", 0),
                line_spacing_pt=s.get("line_spacing_pt", 28),
                line_spacing_rule=s.get("line_spacing_rule", "exactly"))
            render_inline(p, c, s["font_cn"], s["font_en"], s["size_pt"],
                          base_bold=s["bold"])

        prev_type = t

    doc.save(out_path)
    print(f"已生成: {out_path}")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  命令行入口
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description="湘小印·文档出品 - Markdown转Word公文版式工具")
    parser.add_argument("input", help="输入Markdown文件路径")
    parser.add_argument("output", help="输出Word文件路径")
    parser.add_argument("--template", "-t", default="guowei",
                        choices=list(TEMPLATES.keys()),
                        help="版式模板（默认: guowei）")
    args = parser.parse_args()

    md_to_docx(args.input, args.output, template_name=args.template)
